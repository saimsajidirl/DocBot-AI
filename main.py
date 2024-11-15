from fastapi import FastAPI, HTTPException, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List
import motor.motor_asyncio
from bson import ObjectId
from docx import Document
import requests
import json
from haystack import Document as HaystackDocument, Pipeline
from haystack.components.retrievers.in_memory import InMemoryBM25Retriever
from haystack.document_stores.in_memory import InMemoryDocumentStore
from haystack_integrations.components.generators.ollama import OllamaGenerator
from haystack.components.builders.prompt_builder import PromptBuilder

app = FastAPI()

client = motor.motor_asyncio.AsyncIOMotorClient("mongodb://localhost:27017")
db = client.mydatabase
collection = db.items


class Item(BaseModel):
    name: str
    description: str
    price: float

class ItemInDB(Item):
    id: str

document_store = InMemoryDocumentStore()
def query_ollama_model(data: str) -> str:
    url = "http://localhost:11434/v1/chat/completions"
    headers = {"Content-Type": "application/json"}
    payload = {
        "model": "qwen2.5:0.5b",
        "messages": [{"role": "user", "content": data}],
    }

    try:
        response = requests.post(url, headers=headers, data=json.dumps(payload))
        response.raise_for_status()

    except requests.exceptions.RequestException as e:
        raise HTTPException(status_code=500, detail=f"Error querying Ollama model: {str(e)}")

    response_data = response.json()
    answer = response_data.get("choices", [{}])[0].get("message", {}).get("content", "")
    if not answer:
        raise HTTPException(status_code=500, detail="Ollama response does not contain a valid answer")

    return answer

@app.get("/generate_docx_from_query/", response_class=FileResponse)
async def generate_docx_from_query(word: str = Query(..., description="Word to search in the description")):
    try:
        items_cursor = collection.find({"description": {"$regex": word, "$options": "i"}})
        items = await items_cursor.to_list(length=100)

        if not items:
            raise HTTPException(status_code=404, detail="No items found with the given word in description")

        haystack_docs = [HaystackDocument(content=item['description'], id=str(item['_id'])) for item in items]
        document_store.write_documents(haystack_docs)

        prompt_template = """
        Given the following documents, answer the query:
        {% for document in documents %}
            {{ document.content }}
        {% endfor %}

        Question: {{ query }}?
        """

        retriever = InMemoryBM25Retriever(document_store=document_store)
        prompt_builder = PromptBuilder(template=prompt_template)
        ollama_generator = OllamaGenerator(model="qwen2.5:0.5b", url="http://localhost:11434")

        pipe = Pipeline()
        pipe.add_component("retriever", retriever)
        pipe.add_component("prompt_builder", prompt_builder)
        pipe.add_component("llm", ollama_generator)
        pipe.connect("retriever", "prompt_builder.documents")
        pipe.connect("prompt_builder", "llm")

        response = pipe.run({"prompt_builder": {"query": word}, "retriever": {"query": word}})
        ollama_answer = response["llm"]["replies"][0]

        doc = Document()
        doc.add_heading('Ollama Model Response', 0)
        doc.add_paragraph(f"Query: {word}")
        doc.add_paragraph(f"Ollama Model Answer: {ollama_answer}")

        doc_filename = "Ollama_Response.docx"
        doc.save(doc_filename)

        return FileResponse(doc_filename,
                            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                            filename=doc_filename)

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")

@app.get("/")
async def read_root():
    return {
        "message": "Welcome! type /docs at the end of the link to create  documents using query"
    }

@app.get("/items/", response_model=List[ItemInDB])
async def get_items():
    try:
        items_cursor = collection.find()
        items = await items_cursor.to_list(length=100)
        return [ItemInDB(**item) for item in items]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error retrieving items: {str(e)}")

@app.get("/items/{item_id}", response_model=ItemInDB)
async def get_item(item_id: str):
    try:
        item = await collection.find_one({"_id": ObjectId(item_id)})
        if item is None:
            raise HTTPException(status_code=404, detail="Item not found")
        return ItemInDB(**item)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error retrieving item: {str(e)}")

@app.post("/items/", response_model=ItemInDB)
async def create_item(item: Item, custom_id: str = None):
    try:
        item_data = item.model_dump()
        if custom_id:
            item_data["_id"] = custom_id
        else:
            item_data["_id"] = str(ObjectId())

        result = await collection.insert_one(item_data)
        created_item = await collection.find_one({"_id": item_data["_id"]})
        return ItemInDB(**created_item)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating item: {str(e)}")

@app.put("/items/{item_id}", response_model=ItemInDB)
async def update_item(item_id: str, item: Item):
    try:
        updated_item = await collection.find_one_and_update(
            {"_id": ObjectId(item_id)},
            {"$set": item.model_dump()},
            return_document=True
        )
        if updated_item is None:
            raise HTTPException(status_code=404, detail="Item not found")
        return ItemInDB(**updated_item)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error updating item: {str(e)}")

@app.delete("/items/{item_id}", response_model=ItemInDB)
async def delete_item(item_id: str):
    try:
        deleted_item = await collection.find_one_and_delete({"_id": ObjectId(item_id)})
        if deleted_item is None:
            raise HTTPException(status_code=404, detail="Item not found")
        return ItemInDB(**deleted_item)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting item: {str(e)}")